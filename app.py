import streamlit as st
import pandas as pd
import pyodbc
from docx import Document
from docx.shared import Pt
from datetime import date, datetime

def get_database_connection():
    """
    Establishes and returns a connection to the SQL Server database.
    """
    connection_string = """DRIVER={ODBC Driver 17 for SQL Server};SERVER=192.168.13.10;
                           DATABASE=IntegraLive;UID=amog;PWD=Abcd#123;Connection Timeout=30"""
    return pyodbc.connect(connection_string)

def fetch_data_from_database(pms_account_code):
    """
    Fetches data from the database for a single PMS account code.
    """
    conn = get_database_connection()
    cursor = conn.cursor()
    
    # Use parameterized query to avoid SQL injection
    sql_query = """
    SELECT 
        A.client_code,
        A.clientname AS Accountname, 
        C.Clientname, 
        ISNULL(a.address1,'') + ' ' + ISNULL(a.CITY,'') + ' ' + ISNULL(a.State,'') + ' ' + ISNULL(a.pin,'') + ' ' + ISNULL(a.Country,'') AS Address,  
        A.DATE_OF_BIRTH, 
        A.ActiveDate,
        A.backofficecodeequity, 
        D.MainObjective AS SchemeCodename,
        G.TypeDesc AS BenchMark, 
        A.EMAIL,
        A.mobile_no,
        B.Int_Name AS IntroducerDistributorName,
        '' AS FeesCommissionDistributor, 
        A.ctPersonDecision AS RelationshipManager,
        A.ctPersonDEmail AS RMEmail, 
        '' AS MobileNo, 
        E.OpeningEquityCorpus + E.OpeningCashCorpus AS TotalCorpusIntroduced, 
        E.OpeningCashCorpus AS Fund, 
        E.OpeningEquityCorpus AS Securities, 
        A.Usr_clientid AS LoginId, 
        A.PAN_no AS Password
    FROM hdr_client A 
        INNER JOIN HDR_Scheme D ON A.SchemeCode = D.SchemeCode  
        INNER JOIN HDR_ClientHead C ON A.head_clientcode = C.Client_code
        INNER JOIN HDR_Intermediary B ON A.inter_code = B.int_code 
        INNER JOIN hist_clientnav E ON A.client_code = E.ClientCode 
        INNER JOIN dtl_schemeportfolio_benchmark_map F ON D.SchemeCode = F.Scheme_Code
        INNER JOIN HDR_SensexType G ON G.TypeCode = F.Benchmarkindices
    WHERE 
        A.backofficecodeequity = ? 
        AND A.SubBrokerCode IS NOT NULL
        AND E.NavAsOn in (select MAX(NavAsOn) from hist_clientnav);
    """
    
    try:
        df = pd.read_sql_query(sql_query, conn, params=[pms_account_code])
        return df
    finally:
        conn.close()

def generate_welcome_letter(pms_account_code):
    """
    Generates a welcome letter for a given PMS account code.
    """
    data_from_database = fetch_data_from_database(pms_account_code)
    if not data_from_database.empty:
        data_from_database['ActiveDate'] = pd.to_datetime(data_from_database['ActiveDate'])
    
        data = {
            "date" : date.today().strftime('%d/%m/%Y') ,
            "Address": data_from_database.Address.iloc[0],
            "Client Name": data_from_database.Clientname.iloc[0],
            "Date of Activation": data_from_database.ActiveDate.dt.strftime('%d-%m-%Y').iloc[0],
            "PMS Account Code": data_from_database.backofficecodeequity.iloc[0],
            "Strategy Opted": data_from_database.SchemeCodename.iloc[0],
            "Strategy Bench Mark": data_from_database.BenchMark.iloc[0],
            "Registered email id": data_from_database.EMAIL.iloc[0],
            "Registered Mobile no.": "+91"+ data_from_database.mobile_no.iloc[0],
            "Name of Distributor": data_from_database.IntroducerDistributorName.iloc[0],
            "Name of RM": data_from_database.RelationshipManager.iloc[0].upper().split('-')[0].strip(),
            "RM email id": data_from_database.RMEmail.iloc[0],
            "Mobile no.": data_from_database.MobileNo.iloc[0],
            "Total Corpus":"Rs."+ " "+str(data_from_database.TotalCorpusIntroduced.iloc[0]),
            "Fund": "Rs."+ " "+ str(data_from_database.Fund.iloc[0]),
            "Securities": "Rs."+ " "+ str(data_from_database.Securities.iloc[0]),
            "Login Id": data_from_database.LoginId.iloc[0],
            "pass": data_from_database.Password.iloc[0],  
            # Add more placeholders as needed
        }
        
        doc = Document(r"C:\Users\amograne\Welcome Letter Folder\welcome_letter_draft.docx")
        update_placeholders(doc, data)

        output_path = f"welcome_letter_{pms_account_code}.docx"
        doc.save(output_path)
        return output_path
    else:
        return None

# Function to update placeholders in both paragraphs and tables
def update_placeholders(doc, data):
    """
    This Function is designed to search and replace placeholder text within a
    Microsoft Word document using the python-docx library. 
    """
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)  # Set the font size to 9 for each run

        for key, value in data.items():
            placeholder = f"<<{key}>>"
            if placeholder in paragraph.text:
                # Replace in the entire paragraph text
                paragraph.text = paragraph.text.replace(placeholder, str(value))
                for run in paragraph.runs:
                    run.font.size = Pt(9)  # Set the font size after replacement

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f"<<{key}>>"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(9)  # Set the font size after replacement

def fetch_account_codes_by_date(start_date):
    """
    Fetches all the PMS account codes from the specified start date.
    """
    conn = get_database_connection()
    
    # Ensure start_date is a datetime object and format it to a string in YYYY-MM-DD format
    if isinstance(start_date, str):
        # Assuming start_date is in 'YYYY-MM-DD' format; adjust if necessary
        start_date = datetime.strptime(start_date, '%Y-%m-%d')
    start_date_formatted = start_date.strftime('%Y-%m-%d')
    
    # Directly integrate the formatted start_date into the SQL query string to avoid parameter binding issue
    sql_query = f"""SELECT backofficecodeequity FROM hdr_client
                   WHERE ActiveDate >= '{start_date_formatted}' AND ActiveDate <= GETDATE()
                   AND SubBrokerCode IS NOT NULL;"""
    
    try:
        df = pd.read_sql_query(sql_query, conn)
        return df
    finally:
        conn.close()


def generate_welcome_letters_from_date(start_date):
    """
    Generates welcome letters for all PMS accounts activated from a start date.
    """
    if isinstance(start_date, str):
        start_date = datetime.strptime(start_date, '%Y-%m-%d')

    account_codes_df = fetch_account_codes_by_date(start_date)
    for code in account_codes_df['backofficecodeequity']:
        letter_path = generate_welcome_letter(code)
        if letter_path:
            print(f"Welcome letter generated for account code {code} at {letter_path}")
        else:
            print(f"Failed to generate letter for account code {code}. Data might be missing.")

# Example usage:
# generate_welcome_letters_from_date('2024-02-02')

def webapp(start_date):
    """
    UI wrapper for generating welcome letters from a start date, with Streamlit.
    """
    account_codes_df = fetch_account_codes_by_date(start_date)
    success, fail = 0, 0
    for code in account_codes_df['backofficecodeequity']:
        letter_path = generate_welcome_letter(code)
        if letter_path:
            st.success(f"Welcome letter generated for account code {code}")
            success += 1
        else:
            st.error(f"Failed to generate letter for account code {code}. Data might be missing.")
            fail += 1
    st.write(f"Process completed. {success} letters generated successfully, {fail} failed.")

# Streamlit UI
st.title("Welcome Letter Generator")

# Streamlit date input for selecting the start date
start_date = st.date_input("Select a start date for generating welcome letters:", datetime.today())

# Button to trigger the welcome letter generation process
if st.button('Generate Welcome Letters'):
    webapp(start_date)
